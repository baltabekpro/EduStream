import os
import logging
from typing import Optional
from pypdf import PdfReader
from docx import Document
from fastapi import UploadFile
from app.core.config import settings
from app.services.ocr_service import ocr_service

logger = logging.getLogger(__name__)


class FileProcessor:
    """Service for file processing and text extraction."""
    
    async def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    async def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    async def save_file(self, upload_file: UploadFile, material_id: str) -> str:
        """Save uploaded file to disk."""
        upload_dir = settings.UPLOAD_DIR
        os.makedirs(upload_dir, exist_ok=True)
        
        # Create unique filename
        file_extension = os.path.splitext(upload_file.filename)[1]
        destination = os.path.join(upload_dir, f"{material_id}{file_extension}")
        
        with open(destination, "wb") as buffer:
            content = await upload_file.read()
            if len(content) > settings.MAX_UPLOAD_SIZE:
                raise ValueError(f"File is too large. Max size is {settings.MAX_UPLOAD_SIZE} bytes")
            buffer.write(content)
        
        return destination
    
    async def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return text.strip()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    text = f.read()
                return text.strip()
            except Exception as e:
                raise ValueError(f"Failed to extract text from TXT: {str(e)}")
        except Exception as e:
            raise ValueError(f"Failed to extract text from TXT: {str(e)}")
    
    async def extract_text(self, file_path: str) -> str:
        """Extract text from file based on extension."""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        logger.info(f"Extracting text from {file_path} (extension: {file_extension})")
        
        try:
            if file_extension == ".pdf":
                return await self.extract_text_from_pdf(file_path)
            elif file_extension == ".docx":
                return await self.extract_text_from_docx(file_path)
            elif file_extension == ".txt":
                return await self.extract_text_from_txt(file_path)
            elif file_extension in [".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"]:
                return await ocr_service.extract_text_from_image(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}. Supported: .pdf, .docx, .txt, images")
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise


# Legacy functions for backward compatibility
async def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file."""
    try:
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")


async def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    try:
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


async def save_upload_file(upload_file: UploadFile, destination: str) -> str:
    """Save uploaded file to disk."""
    os.makedirs(os.path.dirname(destination), exist_ok=True)
    
    with open(destination, "wb") as buffer:
        content = await upload_file.read()
        buffer.write(content)
    
    return destination


async def extract_text_from_txt(file_path: str) -> str:
    """Extract text from TXT file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return text.strip()
    except UnicodeDecodeError:
        # Try with different encoding
        try:
            with open(file_path, 'r', encoding='latin-1') as f:
                text = f.read()
            return text.strip()
        except Exception as e:
            raise ValueError(f"Failed to extract text from TXT: {str(e)}")
    except Exception as e:
        raise ValueError(f"Failed to extract text from TXT: {str(e)}")


async def process_uploaded_file(upload_file: UploadFile) -> tuple[str, str]:
    """
    Process uploaded file and extract text.
    Returns tuple of (file_path, extracted_text).
    """
    # Validate file extension
    file_extension = os.path.splitext(upload_file.filename)[1].lower()
    if file_extension not in [".pdf", ".docx", ".txt", ".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"]:
        raise ValueError("Only PDF, DOCX, TXT, and image files are supported")
    
    # Save file
    upload_dir = settings.UPLOAD_DIR
    os.makedirs(upload_dir, exist_ok=True)
    
    file_path = os.path.join(upload_dir, upload_file.filename)
    await save_upload_file(upload_file, file_path)
    
    # Extract text
    logger.info(f"Processing file: {file_path}")
    if file_extension == ".pdf":
        text = await extract_text_from_pdf(file_path)
    elif file_extension == ".docx":
        text = await extract_text_from_docx(file_path)
    elif file_extension == ".txt":
        text = await extract_text_from_txt(file_path)
    elif file_extension in [".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"]:
        text = await ocr_service.extract_text_from_image(file_path)
    else:
        raise ValueError("Unsupported file type")
    
    logger.info(f"Extracted {len(text)} characters from {file_path}")
    return file_path, text


# Global instance
file_processor = FileProcessor()
